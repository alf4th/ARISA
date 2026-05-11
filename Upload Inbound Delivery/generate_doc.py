from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────
SAP_BLUE   = RGBColor(0x00, 0x70, 0xF2)
DARK_GRAY  = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
CODE_GRAY  = RGBColor(0x2D, 0x2D, 0x2D)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = SAP_BLUE
    p.runs[0].font.size = Pt(16)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = DARK_GRAY
    p.runs[0].font.size = Pt(13)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = DARK_GRAY
    p.runs[0].font.size = Pt(11)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10) if p.runs else None
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = CODE_GRAY
    shading = OxmlElement('w:pPr')
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def info_table(rows, col_widths=None):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Table Grid'
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        t.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        if i == 0:
            set_cell_bg(t.rows[i].cells[0], '0070F2')
            set_cell_bg(t.rows[i].cells[1], '0070F2')
            t.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            t.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        elif i % 2 == 0:
            set_cell_bg(t.rows[i].cells[0], 'F0F0F0')
            set_cell_bg(t.rows[i].cells[1], 'F0F0F0')
    if col_widths:
        for row in t.rows:
            row.cells[0].width = Cm(col_widths[0])
            row.cells[1].width = Cm(col_widths[1])
    return t

def col_table(headers, data_rows, col_widths=None):
    t = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    t.style = 'Table Grid'
    # header row
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, '0070F2')
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        run.font.size = Pt(9)
    # data rows
    for i, row in enumerate(data_rows):
        for j, val in enumerate(row):
            cell = t.rows[i+1].cells[j]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 1:
                set_cell_bg(cell, 'F5F5F5')
    if col_widths:
        for row in t.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)
    return t

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TECHNICAL DOCUMENTATION')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = SAP_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Upload Inbound Delivery')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = DARK_GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SAP Fiori Application')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

meta = [
    ('Project', 'Upload Inbound Delivery'),
    ('App ID', 'uploadinbounddelivery.uploadinbounddelivery'),
    ('Version', '0.0.1'),
    ('ABAP Object Name', 'ZINBDLB_UPLOAD'),
    ('ABAP Package', 'ZCUST_APP_017'),
    ('SAP System', 'my424750.s4hana.cloud.sap'),
    ('OData Service', 'zui_tbinb_dlv_o4 (V4)'),
    ('UI5 Min Version', '1.136.7'),
    ('SAP Layer', 'CUSTOMER_BASE'),
    ('Date', datetime.date.today().strftime('%B %d, %Y')),
]
t = doc.add_table(rows=len(meta), cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(meta):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    t.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    t.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_bg(t.rows[i].cells[0], 'E8F4FD')
for row in t.rows:
    row.cells[0].width = Cm(6)
    row.cells[1].width = Cm(9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════
heading1('1. Overview')
body('The Upload Inbound Delivery application is a custom SAP Fiori freestyle application '
     'developed for PT ARISA (Customer Layer). It enables warehouse or logistics teams to '
     'upload inbound delivery data from a structured Excel template directly into the '
     'SAP S/4HANA Cloud system via OData V4 API.')
doc.add_paragraph()

heading2('1.1  Purpose')
bullet('Parse and upload inbound delivery line items from an Excel file (.xlsx)')
bullet('Preview and validate data before submitting to the backend')
bullet('Monitor newly uploaded data vs historical data via a tabbed list page')
bullet('Export table data back to Excel for reporting')
bullet('Support for marking records as "cleared" to move them to history')

doc.add_paragraph()
heading2('1.2  Key Features')
col_table(
    ['Feature', 'Description'],
    [
        ['Excel Upload',        'Parse ".xlsx" file using SheetJS library (reads "Inbound template" sheet from row 11)'],
        ['Client Validation',   '8 required fields validated before save; empty rows filtered automatically'],
        ['Batch OData Create',  'All rows submitted in a single OData V4 batch request for performance'],
        ['Dual-Tab List View',  '"Newly Uploaded" (IsCleared=0) and "History" (IsCleared=1) tabs with auto-filter'],
        ['Clear Data',          'Multi-select rows and mark as cleared (IsCleared=1) via batch PATCH'],
        ['Export to Excel',     'Export current tab data using sap.ui.export.Spreadsheet (25 columns)'],
        ['Full-Width Layout',   'Overrides FLP limited-width CSS for full-screen table display'],
        ['Draft Support',       'OData V4 service exposes draft actions (Prepare, Activate, Discard)'],
    ],
    [7, 10]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ══════════════════════════════════════════════════════════════
heading1('2. Application Architecture')

heading2('2.1  Technology Stack')
col_table(
    ['Layer', 'Technology', 'Version / Detail'],
    [
        ['Frontend Framework', 'SAPUI5',              'min 1.136.7, loaded from https://ui5.sap.com'],
        ['UI Pattern',         'Fiori Freestyle',      'MVC pattern (XML Views + JS Controllers)'],
        ['Data Binding',       'OData V4 Model',       'sap.ui.model.odata.v4.ODataModel'],
        ['Excel Parsing',      'SheetJS (XLSX.js)',    'Loaded from CDN + local copy in lib/'],
        ['Export',             'sap.ui.export',        'Spreadsheet class for Excel export'],
        ['Routing',            'sap.m.routing.Router', 'Hash-based navigation'],
        ['Theme',              'SAP Horizon',          'sap_horizon'],
        ['Dev Tooling',        'SAP Fiori Tools',      '@sap/ux-ui5-tooling 1.23.x'],
        ['Backend',            'SAP S/4HANA Cloud',    'my424750.s4hana.cloud.sap'],
        ['Authentication',     'Reentrance Ticket',    'SAML-based SSO via fiori-tools-proxy'],
    ],
    [4, 5, 8]
)

doc.add_paragraph()
heading2('2.2  Application Flow')
body('The following describes the end-to-end user journey through the application:')
doc.add_paragraph()

steps = [
    ('Step 1 — List Page (default route)',
     'User lands on the List Page. The "Newly Uploaded Data" tab is active by default. '
     'The table is filtered to show records where IsCleared = 0, sorted by LastChangedAt descending.'),
    ('Step 2 — Navigate to Upload Page',
     'User clicks "Go to Upload Page" button. Router navigates to the CustomPage route. '
     'Any previously loaded Excel data and file selection are cleared on entry.'),
    ('Step 3 — Select Excel File',
     'User selects an .xlsx file via the FileUploader control. Only .xlsx files are accepted. '
     'The file reference is stored in the controller\'s _file property.'),
    ('Step 4 — Preview Data',
     'User clicks "Preview Data". The controller uses FileReader API + SheetJS to parse the '
     '"Inbound template" sheet starting at row 11. Empty rows are filtered. '
     'Numeric dates and batch numbers are normalized. Required fields are validated. '
     'Errors are displayed via MessageBox. Valid data is bound to the preview table via JSONModel.'),
    ('Step 5 — Save Data',
     'User clicks "Save". The controller maps Excel columns to OData entity properties and '
     'issues a batch OData V4 create (bindList + create) for all rows simultaneously. '
     'A loading indicator is shown. On success, a MessageBox confirms the count and navigates '
     'back to the List Page.'),
    ('Step 6 — View & Manage Data',
     'The List Page refreshes and shows newly created records. User can select records and '
     'click "Clear Data" to move them to the History tab (sets IsCleared = 1 via batch PATCH). '
     'Either tab supports "Export to Excel".'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = SAP_BLUE
    body(desc)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════
heading1('3. Project Structure')

heading2('3.1  Directory Layout')
code_block(
'Upload Inbound Delivery/\n'
'├── package.json                    # NPM config & scripts\n'
'├── ui5.yaml                        # Main dev config (real backend)\n'
'├── ui5-mock.yaml                   # Mock server config\n'
'├── ui5-local.yaml                  # Local dev with offline SAPUI5\n'
'├── ui5-deploy.yaml                 # Deployment config (ABAP)\n'
'└── webapp/\n'
'    ├── index.html                  # Standalone entry point\n'
'    ├── Component.js                # UI5 Component root\n'
'    ├── manifest.json               # App descriptor\n'
'    ├── controller/\n'
'    │   ├── App.controller.js       # Root controller (full-width fix)\n'
'    │   ├── ListPage.controller.js  # List / tab / export / clear logic\n'
'    │   ├── CustomPage.controller.js# Upload / preview / save logic\n'
'    │   └── inbound_delivery.controller.js  # Placeholder\n'
'    ├── view/\n'
'    │   ├── App.view.xml            # Root shell view\n'
'    │   ├── ListPage.view.xml       # Tabbed data list view\n'
'    │   ├── CustomPage.view.xml     # Upload page view\n'
'    │   └── inbound_delivery.view.xml       # Placeholder\n'
'    ├── model/\n'
'    │   └── models.js               # Device model factory\n'
'    ├── i18n/\n'
'    │   └── i18n.properties         # Translatable texts\n'
'    ├── css/\n'
'    │   └── style.css               # Full-width override styles\n'
'    ├── lib/\n'
'    │   └── xlsx.full.min.js        # Local SheetJS copy\n'
'    └── localService/\n'
'        └── mainService/\n'
'            ├── metadata.xml        # OData V4 service metadata\n'
'            └── data/               # Mock data JSON files\n'
)

doc.add_paragraph()
heading2('3.2  NPM Scripts')
col_table(
    ['Script', 'Command', 'Description'],
    [
        ['start',                 'fiori run',                              'Start with real S/4HANA backend (reentrance ticket auth)'],
        ['start-mock',            'fiori run --config ui5-mock.yaml',       'Start with local mock server (no backend needed)'],
        ['start-local',           'fiori run --config ui5-local.yaml',      'Start with offline SAPUI5 + mock server'],
        ['start-noflp',           'fiori run --open /index.html',           'Start without Fiori Launchpad (standalone)'],
        ['build',                 'ui5 build --config=ui5.yaml',            'Build optimized bundle to /dist folder'],
        ['deploy',                'npm run build && fiori deploy',          'Build + deploy to ABAP system'],
        ['deploy-test',           'fiori deploy --testMode true',           'Dry-run deployment (validates config)'],
        ['undeploy',              'fiori undeploy --config ui5-deploy.yaml','Remove app from ABAP system'],
        ['unit-test',             'fiori run --config ui5-mock.yaml',       'Open unit test runner'],
        ['int-test',              'fiori run --config ui5-mock.yaml',       'Open integration (OPA) test runner'],
    ],
    [3.5, 5.5, 8]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. CONFIGURATION FILES
# ══════════════════════════════════════════════════════════════
heading1('4. Configuration Files')

heading2('4.1  ui5.yaml — Main (Real Backend)')
info_table([
    ['Property', 'Value'],
    ['specVersion', '3.1'],
    ['Middleware', 'fiori-tools-proxy, fiori-tools-appreload, fiori-tools-preview'],
    ['Backend URL', 'https://my424750.s4hana.cloud.sap'],
    ['Backend Path', '/sap'],
    ['Auth Type', 'reentranceTicket (SAML SSO)'],
    ['UI5 Source', 'https://ui5.sap.com (latest)'],
    ['Live Reload Port', '35729'],
    ['FLP Theme', 'sap_horizon'],
], [5, 12])

doc.add_paragraph()
heading2('4.2  ui5-mock.yaml — Mock Server')
info_table([
    ['Property', 'Value'],
    ['Extra Middleware', 'sap-fe-mockserver (before CSP)'],
    ['Mock Mount Path', '/'],
    ['OData URL Path', '/sap/opu/odata4/sap/zui_tbinb_dlv_o4/srvd/sap/zui_tbinb_dlv_o4/0001'],
    ['Metadata Source', './webapp/localService/mainService/metadata.xml'],
    ['Mock Data Path', './webapp/localService/mainService/data'],
    ['Generate Mock Data', 'true (auto-generated if data/ is empty)'],
], [5, 12])

doc.add_paragraph()
heading2('4.3  ui5-local.yaml — Offline Development')
info_table([
    ['Property', 'Value'],
    ['Framework', 'SAPUI5 1.136.0 (local install)'],
    ['Libraries', 'sap.m, sap.ui.core, sap.ushell, themelib_sap_horizon'],
    ['Includes', 'Both proxy (real backend) AND mock server'],
    ['Use Case', 'Development without internet access (UI5 served locally)'],
], [5, 12])

doc.add_paragraph()
heading2('4.4  ui5-deploy.yaml — Deployment to ABAP')
info_table([
    ['Property', 'Value'],
    ['Deploy Task', 'deploy-to-abap (after generateCachebusterInfo)'],
    ['Target URL', 'https://my424750.s4hana.cloud.sap'],
    ['Auth Type', 'reentranceTicket'],
    ['ABAP App Name', 'ZINBDLB_UPLOAD'],
    ['Description', 'Inbound Delivery Upload'],
    ['ABAP Package', 'ZCUST_APP_017'],
    ['Transport Request', 'Read from environment variable: TRANSPORT_REQUEST'],
    ['Excluded from Build', '/test/, /localService/'],
], [5, 12])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. MANIFEST
# ══════════════════════════════════════════════════════════════
heading1('5. Application Manifest (manifest.json)')

heading2('5.1  Identification')
info_table([
    ['Property', 'Value'],
    ['App ID', 'uploadinbounddelivery.uploadinbounddelivery'],
    ['Version', '0.0.1'],
    ['Min UI5 Version', '1.136.7'],
    ['Semantic Object', 'InboudDeliverySemanticObject'],
    ['FLP Action', 'manage'],
    ['FLP Title', 'Upload Inbound Delivery'],
], [5, 12])

doc.add_paragraph()
heading2('5.2  OData Data Source')
info_table([
    ['Property', 'Value'],
    ['Data Source Name', 'mainService'],
    ['Service URI', '/sap/opu/odata4/sap/zui_tbinb_dlv_o4/srvd/sap/zui_tbinb_dlv_o4/0001/'],
    ['Type', 'OData V4'],
    ['Local Metadata', 'localService/mainService/metadata.xml'],
    ['Operation Mode', 'Server (server-side filtering/sorting)'],
    ['Auto Expand Select', 'true'],
    ['Early Requests', 'true'],
], [5, 12])

doc.add_paragraph()
heading2('5.3  UI5 Libraries')
col_table(
    ['Library', 'Purpose'],
    [
        ['sap.m',          'Mobile-first controls (Page, Button, IconTabBar, MessageBox, etc.)'],
        ['sap.ui.core',    'Core framework, MVC, routing, models'],
        ['sap.ui.comp',    'Smart controls (SmartFilterBar, SmartTable)'],
        ['sap.ui.table',   'Grid Table for large data sets (odataTable, historyTable, previewTable)'],
        ['sap.ui.unified', 'FileUploader control'],
    ],
    [5, 12]
)

doc.add_paragraph()
heading2('5.4  Routing')
col_table(
    ['Route Name', 'Pattern', 'Target View', 'Notes'],
    [
        ['RouteListPage',          '(empty)',    'ListPage',          'Default route — landing page'],
        ['Routeinbound_delivery',  ':?query:',   'inbound_delivery',  'Placeholder, not used in current flow'],
        ['CustomPage',             'custom',     'CustomPage',        'Upload page'],
    ],
    [5, 3.5, 4, 4.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. CONTROLLERS
# ══════════════════════════════════════════════════════════════
heading1('6. Controllers')

heading2('6.1  App.controller.js')
body('Root controller attached to the App shell view.')
bullet('Removes the CSS class sapUShellApplicationContainerLimitedWidth from the FLP shell '
       'container on initialization, enabling the application to use the full browser width.')

doc.add_paragraph()
heading2('6.2  CustomPage.controller.js  (Upload Page)')
body('Handles all Excel upload, parse, validate, and save logic.')
doc.add_paragraph()

heading3('onInit')
bullet('Creates a JSONModel with empty excelData array and sets it on the view.')
bullet('Attaches _onRouteMatched handler to the "CustomPage" route.')

heading3('_onRouteMatched / _clearData')
bullet('Clears _file reference, resets /excelData to empty array.')
bullet('Calls fileUploader.clear() and fileUploader.setValue("") to reset the file picker UI.')

heading3('onFileChange')
bullet('Stores the selected File object in this._file for later reading.')

heading3('onPreview  (main parse logic)')
bullet('Validates that a file is selected.')
bullet('Uses FileReader.readAsArrayBuffer → XLSX.read to parse the workbook.')
bullet('Reads the sheet named "Inbound template" (case-sensitive). Shows error if not found.')
bullet('Calls XLSX.utils.sheet_to_json with range:10 (skips first 10 rows — header rows) and fixed column headers.')
bullet('Filters out empty rows (all cells blank after trim).')
bullet('Normalizes Delivery Date and Delivery Time: converts numeric Excel serial values to integer strings.')
bullet('Normalizes Batch Number and Purchase Order Item: strips trailing decimal (.0) from numeric values.')
bullet('Validates 8 required fields per row; collects all errors and shows them in a single MessageBox.error.')
bullet('Sets normalized data to /excelData on the JSONModel, binding it to the preview table.')

heading3('Required Fields for Validation')
col_table(
    ['#', 'Excel Column Name', 'OData Property'],
    [
        ['1', 'Pallet Number',         'PalletNumber'],
        ['2', 'Delivery Date',         'DeliveryDate'],
        ['3', 'Delivery Time',         'DeliveryTime'],
        ['4', 'Quantity',              'ActualDeliveryQuantity'],
        ['5', 'Arisa Part Number',     'Material'],
        ['6', 'Plant',                 'Plant'],
        ['7', 'Purchase Order Number', 'PurchaseOrderNumber'],
        ['8', 'Purchase Order Item',   'PurchaseOrderItem'],
    ],
    [1.5, 6.5, 6]
)
doc.add_paragraph()

heading3('onSave  (OData batch create)')
bullet('Maps /excelData items to OData payload objects (21 fields).')
bullet('Sets ActualDeliveryQuantity as a 4-decimal string (toFixed(4)).')
bullet('Sets Response field to "-" as default.')
bullet('Creates a list binding on /ZC_TBINB_DLV and calls binding.create(payload) for each row.')
bullet('Uses Promise.all(contexts.map(c => c.created())) to await all create operations.')
bullet('On success: hides BusyIndicator, shows success MessageBox with record count, navigates to RouteListPage.')
bullet('On error: hides BusyIndicator, shows detailed error MessageBox including error.details array if present.')

doc.add_paragraph()
heading2('6.3  ListPage.controller.js  (Data List Page)')
body('Handles tab filtering, data refresh, clear data, navigation, and Excel export.')
doc.add_paragraph()

heading3('onInit')
bullet('Attaches _onRouteMatched to RouteListPage route.')

heading3('_onRouteMatched')
bullet('Refreshes the OData model and the odataTable row binding.')
bullet('Applies a descending sort on LastChangedAt.')
bullet('Calls _applyFilterToNewlyUploadedTab to set IsCleared EQ 0 filter on the table.')

heading3('onTabSelect')
bullet('"newly" tab: filters odataTable with IsCleared EQ 0.')
bullet('"history" tab: filters historyTable with IsCleared EQ 1.')
bullet('Both tabs apply descending LastChangedAt sort.')

heading3('onClearData')
bullet('Reads selected row indices from the grid table.')
bullet('Shows warning if no rows are selected.')
bullet('Confirms action with MessageBox.confirm showing selected row count.')
bullet('On confirm, calls _clearSelectedData.')

heading3('_clearSelectedData')
bullet('Detects OData V4 (isA check) and uses context.setProperty("IsCleared", 1) + model.submitBatch("updateGroup").')
bullet('Fallback OData V2 path uses model.setProperty + model.submitChanges with batch group.')
bullet('Refreshes binding and clears selection on success.')

heading3('onExport')
bullet('Gets the currently active table via _getCurrentTable (based on IconTabBar selected key).')
bullet('Reads all row contexts from the binding and extracts plain objects.')
bullet('Builds a 25-column Spreadsheet configuration (sap.ui.export.Spreadsheet).')
bullet('Exports file named InboundDelivery_{timestamp}.xlsx.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. VIEWS
# ══════════════════════════════════════════════════════════════
heading1('7. Views')

heading2('7.1  App.view.xml')
body('Root view. Contains a single <App> container (sap.m.App) that hosts all page views '
     'as managed by the router. No visible content of its own.')

doc.add_paragraph()
heading2('7.2  ListPage.view.xml')
body('Main data list view. Contains:')
bullet('A "Go to Upload Page" button (top-right) — navigates to CustomPage route.')
bullet('An IconTabBar with two tabs:')
bullet('"Newly Uploaded Data" (key: newly) — shows odataTable bound to /ZC_TBINB_DLV with IsCleared=0 filter.', level=1)
bullet('"History of Uploaded Data" (key: history) — shows historyTable bound to the same entity set with IsCleared=1 filter.', level=1)
bullet('Each tab has a toolbar with relevant action buttons and a sap.ui.table.Table with 25 columns.')
doc.add_paragraph()

heading2('7.3  Table Columns — ListPage')
col_table(
    ['#', 'Column Label', 'OData Property', 'Notes'],
    [
        ['1',  'Model',                    'Model',                   ''],
        ['2',  'Model Description',        'ModelDescription',        ''],
        ['3',  'SO Number / Line Item',    'SoNumberLineItem',        ''],
        ['4',  'Container Number',         'ContainerNumber',         ''],
        ['5',  'Proforma Invoice Number',  'ProformaInvoiceNumber',   ''],
        ['6',  'Purchase Order Number',    'PurchaseOrderNumber',     ''],
        ['7',  'Pallet Number',            'PalletNumber',            ''],
        ['8',  'Gross Weight',             'GrossWeight',             ''],
        ['9',  'PO Item',                  'PurchaseOrderItem',       ''],
        ['10', 'Arisa Part Number',        'Material',                ''],
        ['11', 'Vendor Part Number',       'VendorPartNumber',        ''],
        ['12', 'Quantity',                 'ActualDeliveryQuantity',  'Decimal 10,4'],
        ['13', 'Delivery Date',            'DeliveryDate',            ''],
        ['14', 'Delivery Time',            'DeliveryTime',            ''],
        ['15', 'Plant',                    'Plant',                   ''],
        ['16', 'Batch Number',             'BatchNumber',             ''],
        ['17', 'Free Goods Indicator',     'FreeGoodsIndicator',      ''],
        ['18', 'Note of Free Goods',       'NoteOfFreeGoods',         ''],
        ['19', 'PO Ref for Free Goods',    'PoRefForFreeGoods',       ''],
        ['20', 'PO Line Item Ref',         'PoLineItemRef',           ''],
        ['21', 'Response',                 'Response',                'Default: "-"'],
        ['22', 'Created By',               'CreatedBy',               ''],
        ['23', 'Changed By',               'ChangedBy',               ''],
        ['24', 'Local Last Changed At',    'LocalLastChangedAt',      ''],
        ['25', 'Last Changed At',          'LastChangedAt',           'Sort field (desc)'],
    ],
    [1, 5, 5, 6]
)

doc.add_paragraph()
heading2('7.4  CustomPage.view.xml')
body('Upload page view. Contains:')
bullet('FileUploader (sap.ui.unified.FileUploader) — accepts .xlsx only.')
bullet('"Preview Data" button — triggers Excel parsing and table population.')
bullet('"Save" button — triggers OData batch create.')
bullet('A sap.ui.table.Table (previewTable) bound to /excelData on the local JSONModel, showing 21 columns.')

doc.add_paragraph()
heading2('7.5  Preview Table Columns — CustomPage')
col_table(
    ['#', 'Column Label', 'Excel Header'],
    [
        ['1',  'No.',                              '(auto-generated row index)'],
        ['2',  'Model',                            'Model'],
        ['3',  'Model Description',                'Model Description'],
        ['4',  'SO Number / Line Item',            'SO Number / Line Item'],
        ['5',  'Container Number',                 'Container Number'],
        ['6',  'Proforma Invoice',                 'Proforma Invoice'],
        ['7',  'Bill of Lading (optional)',         'Bill of Lading (optional)'],
        ['8',  'Purchase Order Number',            'Purchase Order Number'],
        ['9',  'Pallet Number',                    'Pallet Number'],
        ['10', 'Gross Weight',                     'Gross Weight'],
        ['11', 'Purchase Order Item',              'Purchase Order Item'],
        ['12', 'Arisa Part Number',                'Arisa Part Number'],
        ['13', 'Vendor Part Number',               'Vendor Part Number'],
        ['14', 'Quantity',                         'Quantity'],
        ['15', 'Delivery Date',                    'Delivery Date'],
        ['16', 'Delivery Time',                    'Delivery Time'],
        ['17', 'Plant',                            'Plant'],
        ['18', 'Batch Number',                     'Batch Number'],
        ['19', 'Free Goods Indicator',             'Free Goods Indicator'],
        ['20', 'Note of Free Goods',               'Note of Free Goods'],
        ['21', 'PO Ref for Free Goods',            'PO Ref for Free Goods'],
        ['22', 'PO Line Item Ref for Free Goods',  'PO Line item Ref for Free Goods'],
    ],
    [1, 6, 10]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. ODATA SERVICE
# ══════════════════════════════════════════════════════════════
heading1('8. OData V4 Service')

heading2('8.1  Service Details')
info_table([
    ['Property', 'Value'],
    ['Service Name', 'zui_tbinb_dlv_o4'],
    ['Service URL', '/sap/opu/odata4/sap/zui_tbinb_dlv_o4/srvd/sap/zui_tbinb_dlv_o4/0001/'],
    ['Entity Set', 'ZC_TBINB_DLV'],
    ['Entity Type', 'ZC_TBINB_DLVType'],
    ['Key Fields', 'ID (Edm.Guid), IsActiveEntity (Edm.Boolean)'],
    ['Draft Enabled', 'Yes — Prepare, Activate, Discard, Edit, Resume actions'],
    ['Concurrency Token', 'LocalLastChangedAt'],
], [5, 12])

doc.add_paragraph()
heading2('8.2  Entity Properties')
col_table(
    ['Property Name', 'EDM Type', 'Nullable', 'Notes'],
    [
        ['ID',                          'Edm.Guid',     'false', 'Primary key'],
        ['IsActiveEntity',              'Edm.Boolean',  'false', 'Draft key'],
        ['InboundDeliverySequence',     'Edm.String',   'true',  'Auto-generated sequence'],
        ['Model',                       'Edm.String',   'true',  ''],
        ['ModelDescription',            'Edm.String',   'true',  ''],
        ['SoNumberLineItem',            'Edm.String',   'true',  'SO Number / Line Item'],
        ['ContainerNumber',             'Edm.String',   'true',  ''],
        ['ProformaInvoiceNumber',       'Edm.String',   'true',  ''],
        ['BillOfPayment',               'Edm.String',   'true',  'Bill of Lading'],
        ['PurchaseOrderNumber',         'Edm.String',   'true',  ''],
        ['PurchaseOrderItem',           'Edm.String',   'true',  ''],
        ['Material',                    'Edm.String',   'true',  'Arisa Part Number'],
        ['ArisaPartNumber',             'Edm.String',   'true',  ''],
        ['VendorPartNumber',            'Edm.String',   'true',  ''],
        ['ActualDeliveryQuantity',      'Edm.Decimal',  'true',  'Precision 10, Scale 4'],
        ['PalletNumber',                'Edm.String',   'true',  ''],
        ['GrossWeight',                 'Edm.String',   'true',  ''],
        ['DeliveryDate',                'Edm.String',   'true',  ''],
        ['DeliveryTime',                'Edm.String',   'true',  ''],
        ['Plant',                       'Edm.String',   'true',  ''],
        ['BatchNumber',                 'Edm.String',   'true',  ''],
        ['FreeGoodsIndicator',          'Edm.String',   'true',  ''],
        ['NoteOfFreeGoods',             'Edm.String',   'true',  ''],
        ['PoRefForFreeGoods',           'Edm.String',   'true',  ''],
        ['PoLineItemRef',               'Edm.String',   'true',  ''],
        ['Response',                    'Edm.String',   'true',  'Default "-", updated by backend'],
        ['IsCleared',                   'Edm.Int32',    'true',  '0=New, 1=History'],
        ['CreatedBy',                   'Edm.String',   'true',  'Auto-set by backend'],
        ['ChangedBy',                   'Edm.String',   'true',  'Auto-set by backend'],
        ['LocalLastChangedAt',          'Edm.DateTimeOffset', 'true', 'Concurrency token'],
        ['LastChangedAt',               'Edm.DateTimeOffset', 'true', 'Used for default sort'],
    ],
    [5.5, 4, 2.5, 5]
)

doc.add_paragraph()
heading2('8.3  Supported Operations')
col_table(
    ['Operation', 'HTTP Method', 'Used By', 'Notes'],
    [
        ['Create',   'POST',   'CustomPage.onSave',       'Batch create all rows in one request'],
        ['Read',     'GET',    'ListPage (both tabs)',     'Server-side filter + sort'],
        ['Update',   'PATCH',  'ListPage.onClearData',    'Sets IsCleared=1 via batch PATCH'],
        ['Delete',   'DELETE', '(not used in UI)',        'Available in service'],
        ['Activate', 'POST',   '(draft action)',          'Activates draft entity'],
        ['Discard',  'POST',   '(draft action)',          'Discards draft changes'],
        ['Prepare',  'POST',   '(draft action)',          'Validates before activation'],
    ],
    [3, 3, 5, 6]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. EXCEL TEMPLATE
# ══════════════════════════════════════════════════════════════
heading1('9. Excel Template Specification')

heading2('9.1  Template Requirements')
info_table([
    ['Property', 'Value'],
    ['File Format', '.xlsx (Excel 2007+)'],
    ['Required Sheet Name', 'Inbound template  (case-sensitive)'],
    ['Data Start Row', 'Row 11 (rows 1–10 are header/title rows, skipped by parser)'],
    ['First Data Row', 'Row 12 (row 11 is the column header row read by parser as first record, then shifted off)'],
], [5, 12])

doc.add_paragraph()
heading2('9.2  Column Mapping')
col_table(
    ['Col #', 'Excel Column Header', 'Required?', 'OData Property', 'Notes'],
    [
        ['1',  'Model',                              'No',  'Model',                    ''],
        ['2',  'Model Description',                  'No',  'ModelDescription',         ''],
        ['3',  'SO Number / Line Item',              'No',  'SoNumberLineItem',         ''],
        ['4',  'Container Number',                   'No',  'ContainerNumber',          ''],
        ['5',  'Proforma Invoice',                   'No',  'ProformaInvoiceNumber',    ''],
        ['6',  'Bill of Lading (optional)',           'No',  'BillOfPayment',            ''],
        ['7',  'Purchase Order Number',              'Yes', 'PurchaseOrderNumber',      ''],
        ['8',  'Pallet Number',                      'Yes', 'PalletNumber',             ''],
        ['9',  'Gross Weight',                       'No',  'GrossWeight',              ''],
        ['10', 'Purchase Order Item',                'Yes', 'PurchaseOrderItem',        'Decimal suffix stripped'],
        ['11', 'Arisa Part Number',                  'Yes', 'Material',                 ''],
        ['12', 'Vendor Part Number',                 'No',  'VendorPartNumber',         ''],
        ['13', 'Quantity',                           'Yes', 'ActualDeliveryQuantity',   'Formatted as 4-decimal string'],
        ['14', 'Delivery Date',                      'Yes', 'DeliveryDate',             'Excel serial → integer string'],
        ['15', 'Delivery Time',                      'Yes', 'DeliveryTime',             'Excel serial → integer string'],
        ['16', 'Plant',                              'Yes', 'Plant',                    ''],
        ['17', 'Batch Number',                       'No',  'BatchNumber',              'Decimal suffix stripped'],
        ['18', 'Free Goods Indicator',               'No',  'FreeGoodsIndicator',       ''],
        ['19', 'Note of Free Goods',                 'No',  'NoteOfFreeGoods',          ''],
        ['20', 'PO Ref for Free Goods',              'No',  'PoRefForFreeGoods',        ''],
        ['21', 'PO Line item Ref for Free Goods',    'No',  'PoLineItemRef',            ''],
    ],
    [1.5, 5.5, 2, 4.5, 3.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. DEPLOYMENT
# ══════════════════════════════════════════════════════════════
heading1('10. Deployment')

heading2('10.1  Prerequisites')
bullet('@sap/ux-ui5-tooling installed (devDependency)')
bullet('Authenticated to the target S/4HANA Cloud system (reentrance ticket or stored credentials)')
bullet('TRANSPORT_REQUEST environment variable set to a valid transport request number')
bullet('Target ABAP package ZCUST_APP_017 must exist')

doc.add_paragraph()
heading2('10.2  Deployment Steps')

steps_deploy = [
    ('1. Set transport request',   'set TRANSPORT_REQUEST=<your-transport-number>'),
    ('2. Build & deploy',          'npm run deploy'),
    ('3. Dry-run (optional)',       'npm run deploy-test'),
    ('4. Undeploy if needed',      'npm run undeploy'),
]
for label, cmd in steps_deploy:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}:  ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(cmd)
    run2.font.name = 'Courier New'
    run2.font.size = Pt(9)

doc.add_paragraph()
heading2('10.3  Deployment Target Details')
info_table([
    ['Property', 'Value'],
    ['System URL', 'https://my424750.s4hana.cloud.sap'],
    ['App Name (BSP)', 'ZINBDLB_UPLOAD'],
    ['App Description', 'Inbound Delivery Upload'],
    ['Package', 'ZCUST_APP_017'],
    ['Transport', '$TRANSPORT_REQUEST environment variable'],
    ['Excluded Artifacts', '/test/ and /localService/ directories'],
    ['Auth Method', 'reentranceTicket'],
], [5, 12])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 11. DEVELOPMENT GUIDE
# ══════════════════════════════════════════════════════════════
heading1('11. Development Guide')

heading2('11.1  Local Setup')
col_table(
    ['Step', 'Command', 'Notes'],
    [
        ['Clone / open project', '(open in VS Code)',        'Requires SAP Fiori Tools extension'],
        ['Install dependencies', 'npm install',              'Installs all devDependencies'],
        ['Run with real backend','npm run start',            'Requires VPN + reentrance ticket auth'],
        ['Run with mock data',   'npm run start-mock',       'No backend needed'],
        ['Run offline (local UI5)','npm run start-local',    'Needs SAPUI5 installed locally'],
    ],
    [4, 4.5, 8.5]
)

doc.add_paragraph()
heading2('11.2  Authentication (npm run start)')
body('When running against the real backend, the first request triggers a reentrance ticket '
     'authentication flow:')
bullet('A browser tab opens to the SAP S/4HANA system login page.')
bullet('After successful login, the tab shows "You can close this tab now."')
bullet('The proxy stores the ticket and forwards it for all subsequent /sap requests.')
bullet('If a Basic Auth dialog appears, ensure @sap/ux-ui5-tooling is version 1.22.0 or higher '
       '(run npm update @sap/ux-ui5-tooling to update).')

doc.add_paragraph()
heading2('11.3  Adding New Columns')
body('To add a new field from the OData service to the tables:')
bullet('1. Add the column definition to ListPage.view.xml inside the relevant <table:Table>.')
bullet('2. Add the corresponding entry to the aCols array in ListPage.controller.js onExport().')
bullet('3. Add the column to CustomPage.view.xml previewTable and to the column header array in onPreview().')
bullet('4. Map the Excel column to the OData property in the aPayloads mapping inside onSave().')

doc.add_paragraph()
heading2('11.4  Changing the Excel Template Row Offset')
body('The parser reads the "Inbound template" sheet starting at row 11 (range:10 in XLSX options, '
     'zero-indexed). To change this, update the range value in CustomPage.controller.js onPreview():')
code_block('var excelData = XLSX.utils.sheet_to_json(sheet, {\n'
           '  header: [ ... ],\n'
           '  defval: "",\n'
           '  range: 10   // <-- change this (0-indexed row offset)\n'
           '});')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. KNOWN ISSUES & NOTES
# ══════════════════════════════════════════════════════════════
heading1('12. Known Issues & Notes')

col_table(
    ['ID', 'Issue', 'Status', 'Workaround / Note'],
    [
        ['N-01', 'Basic Auth dialog on npm run start',
         'Fixed',
         'Caused by @sap/ux-ui5-tooling < 1.22. Update via: npm update @sap/ux-ui5-tooling'],
        ['N-02', 'inbound_delivery view/controller',
         'Placeholder',
         'Controller and view exist but have no implementation. Not used in current routing flow.'],
        ['N-03', 'SheetJS loaded from CDN in index.html',
         'Note',
         'A local copy also exists in lib/xlsx.full.min.js. CDN version in index.html takes precedence.'],
        ['N-04', 'IsCleared field — Clear Data',
         'Note',
         'Clear Data sets IsCleared=1. There is no "undo clear" feature in the current UI.'],
        ['N-05', 'Draft support',
         'Note',
         'The OData service is draft-enabled, but the frontend uses direct activation. '
         'Draft lifecycle (Prepare/Activate/Discard) is not explicitly managed in the UI.'],
    ],
    [1.5, 5, 2.5, 8]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 13. CHANGE LOG
# ══════════════════════════════════════════════════════════════
heading1('13. Change Log')
col_table(
    ['Version', 'Date', 'Description'],
    [
        ['0.0.1', datetime.date.today().strftime('%Y-%m-%d'), 'Initial release — Upload Inbound Delivery application'],
    ],
    [2.5, 3.5, 11]
)

# ── Save ──────────────────────────────────────────────────────
out_path = r'd:\JOB ISM\2026\ARISA\Upload Inbound Delivery\Upload_Inbound_Delivery_Technical_Documentation.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
